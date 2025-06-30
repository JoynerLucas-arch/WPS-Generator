import re
from enum import Enum, unique
from docx import Document
import labels


def is_no_content_point(p_d):
    p_t = p_d['type']
    if p_t in TemplateAnalyzer.insert_point_no_content_types:
        return True
    return False


class TemplateAnalyzer:
    _content_label_re = re.compile(r'{{(.*?)}}')

    registered_labels = {label.get_type(): label for label in labels.LabelManager.get_labels()}

    insert_point_content_types = [label.get_type() for label in registered_labels.values() if label.has_content()]
    insert_point_no_content_types = [label.get_type() for label in registered_labels.values() if
                                     not label.has_content()]
    insert_point_types = insert_point_no_content_types + insert_point_content_types

    static_datas = {}

    @classmethod
    def update_labels_info(cls):
        cls.registered_labels = {label.get_type(): label for label in labels.LabelManager.get_labels()}
        cls.insert_point_content_types = [label.get_type() for label in cls.registered_labels.values() if
                                          label.has_content()]
        cls.insert_point_no_content_types = [label.get_type() for label in cls.registered_labels.values() if
                                             not label.has_content()]
        cls.insert_point_types = cls.insert_point_no_content_types + cls.insert_point_content_types

    @classmethod
    def register_static_datas(cls):
        cls.static_datas.clear()
        for label in cls.registered_labels.values():
            label.register_static_datas(cls.static_datas)

    @unique
    class CheckCode(Enum):
        SUCCESS = 1
        NAME_REPEAT = -1
        LABEL_FORMAT_ERROR = -2

        def __str__(self):
            return f'{super(TemplateAnalyzer.CheckCode, self).__str__()} {self.value}'

        def is_error(self):
            return self.value < 0

    @classmethod
    def check_template(cls, file_path: str, insert_operation: callable = is_no_content_point) -> dict:
        
        document = Document(file_path)
        insert_points = {}

        # 处理正文部分
        for element in document.element.body:
            if element.tag.endswith('p'):
                paragraph = document.paragraphs[
                    len([e for e in document.element.body[:list(document.element.body).index(element)] if
                         e.tag.endswith('p')])]
                cls._process_paragraph(paragraph, insert_points, insert_operation, document)
            elif element.tag.endswith('tbl'):
                table = document.tables[
                    len([e for e in document.element.body[:list(document.element.body).index(element)] if
                         e.tag.endswith('tbl')])]
                cls._process_table(table, insert_points, insert_operation, document)

        # 处理页眉页脚
        for section in document.sections:
            for part in [section.header, section.footer]:
                if part:
                    for paragraph in part.paragraphs:
                        cls._process_paragraph(paragraph, insert_points, insert_operation, document)
                    for table in part.tables:
                        cls._process_table(table, insert_points, insert_operation, document)
        
        # 扫描所有可能被遗漏的表格和标签
        cls._scan_all_tables(document, insert_points, insert_operation)
        
        # 添加调试信息 - 标签识别结果摘要
        # print(f"\n调试信息 - 标签识别结果摘要:")
        image_labels = []
        for point_name, point_data in insert_points.items():
            if isinstance(point_data, list):
                for pd in point_data:
                    if pd['type'] == 'image':
                        image_labels.append((point_name, pd['text']))
            elif point_data['type'] == 'image':
                image_labels.append((point_name, point_data['text']))

        return {
            "code": cls.CheckCode.SUCCESS,
            "msg": "successful",
            "data": {
                "document": document,
                "insert_points": insert_points
            }
        }

    @classmethod
    def _scan_all_tables(cls, document, insert_points, insert_operation):
        """全文扫描所有可能的表格，包括在复杂结构中的表格，确保不会遗漏任何表格中的标签"""
        # 使用XPath查找文档中的所有表格
        doc_element = document._element
        all_tables_elements = doc_element.xpath('.//w:tbl')
        
        # 获取已经处理过的表格元素，避免重复处理
        processed_table_elements = set()
        for point_data in insert_points.values():
            if isinstance(point_data, list):
                for pd in point_data:
                    if 'table' in pd:
                        processed_table_elements.add(pd['table']._element)
            else:
                if 'table' in point_data:
                    processed_table_elements.add(point_data['table']._element)
        
        # 处理文档中的所有表格
        for table_element in all_tables_elements:
            # 如果已经处理过这个表格元素，跳过
            if table_element in processed_table_elements:
                continue
            
            try:
                # 尝试为表格元素创建一个Table对象
                from docx.table import Table
                table = Table(table_element, document._body)
                cls._process_table(table, insert_points, insert_operation, document)
                # 添加到已处理的表格集合中
                processed_table_elements.add(table_element)
            except Exception as e:
                print(f"警告：处理表格时出错 - {str(e)}")

    @classmethod
    def _process_table(cls, table, insert_points, insert_operation, document):
        """处理表格中的内容标签，增强识别能力"""
        # 递归处理嵌套表格的函数
        def process_nested_tables(cell):
            # 处理单元格中的嵌套表格
            for nested_table in cell.tables:
                cls._process_table(nested_table, insert_points, insert_operation, document)
        
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                # 首先处理单元格中的嵌套表格
                process_nested_tables(cell)
                
                # 然后处理单元格中的段落
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue

                    # 特殊处理：合并所有runs的文本，确保完整捕获标签
                    full_text = paragraph.text
                    matches = cls._content_label_re.finditer(full_text)

                    for match in matches:
                        point = match.group(1)
                        point_split = point.split(':')
                        if len(point_split) != 2:
                            print(f"调试 - 标签格式错误: '{match.group(0)}', 需要形如 '{{{{类型:名称}}}}'")
                            continue

                        point_type, point_name = point_split
                        if point_type not in cls.insert_point_types:
                            print(f"调试 - 标签类型不支持: '{point_type}', 支持的类型有: {cls.insert_point_types}")
                            continue
                            
                        print(f"调试 - 在表格中找到标签: 类型='{point_type}', 名称='{point_name}', 文本='{match.group(0)}'")

                        # 创建包含完整标签的point_data
                        point_data = {
                            'name': point_name,
                            'type': point_type,
                            'text': match.group(0),
                            'run': paragraph.runs[0] if paragraph.runs else None,
                            'run_index': 0,
                            'paragraph': paragraph,
                            'document': document,
                            'start_pos': match.start(),
                            'end_pos': match.end(),
                            'cell': cell,  # 添加单元格引用
                            'table': table,  # 添加表格引用
                            'row_index': row_index,  # 使用枚举索引而不是table.rows.index(row)
                            'cell_index': cell_index  # 使用枚举索引而不是row.cells.index(cell)
                        }

                        # 获取所有相关runs
                        start_pos = match.start()
                        end_pos = match.end()
                        current_pos = 0
                        containing_runs = []

                        for run_idx, run in enumerate(paragraph.runs):
                            run_length = len(run.text)
                            run_start = current_pos
                            run_end = current_pos + run_length

                            if (run_start <= start_pos < run_end or
                                    run_start < end_pos <= run_end or
                                    (start_pos <= run_start and run_end <= end_pos)):
                                containing_runs.append({
                                    'run': run,
                                    'run_index': run_idx,
                                    'start': run_start,
                                    'end': run_end
                                })

                            current_pos += run_length

                        if containing_runs:
                            point_data['containing_runs'] = containing_runs

                        # 处理标签
                        if not insert_operation(point_data):
                            if point_name in insert_points:
                                if isinstance(insert_points[point_name], list):
                                    insert_points[point_name].append(point_data)
                                else:
                                    insert_points[point_name] = [insert_points[point_name], point_data]
                            else:
                                insert_points[point_name] = point_data

    @classmethod
    def _process_paragraph(cls, paragraph, insert_points, insert_operation, document):
        # 构建run的映射关系
        run_text_map = []
        total_offset = 0
        for run_index, run in enumerate(paragraph.runs):
            run_length = len(run.text)
            run_text_map.append({
                'run': run,
                'run_index': run_index,
                'start': total_offset,
                'end': total_offset + run_length
            })
            total_offset += run_length

        paragraph_text = paragraph.text
        matches = cls._content_label_re.finditer(paragraph_text)

        for match in matches:
            point = match.group(1)
            point_split = point.split(':')
            if len(point_split) != 2:
                print(f"调试 - 标签格式错误: '{match.group(0)}', 需要形如 '{{{{类型:名称}}}}'")
                continue

            point_type, point_name = point_split
            if point_type not in cls.insert_point_types:
                print(f"调试 - 标签类型不支持: '{point_type}', 支持的类型有: {cls.insert_point_types}")
                continue
                
            print(f"调试 - 在段落中找到标签: 类型='{point_type}', 名称='{point_name}', 文本='{match.group(0)}'")

            # 找到标记所在的run
            start_pos = match.start()
            end_pos = match.end()
            containing_runs = []
            for run_info in run_text_map:
                if (run_info['start'] <= start_pos < run_info['end'] or
                        run_info['start'] < end_pos <= run_info['end'] or
                        (start_pos <= run_info['start'] and run_info['end'] <= end_pos)):
                    containing_runs.append(run_info)

            if not containing_runs:
                continue

            point_data = {
                'name': point_name,
                'type': point_type,
                'text': match.group(0),
                'run': containing_runs[0]['run'],
                'run_index': containing_runs[0]['run_index'],
                'paragraph': paragraph,
                'document': document,
                'containing_runs': containing_runs,
                'start_pos': start_pos,
                'end_pos': end_pos
            }

            if not insert_operation(point_data):
                if point_name in insert_points:
                    if isinstance(insert_points[point_name], list):
                        insert_points[point_name].append(point_data)
                    else:
                        insert_points[point_name] = [insert_points[point_name], point_data]
                else:
                    insert_points[point_name] = point_data

    @staticmethod
    def print_check_info(check_info: dict, show_detail=False):
        if check_info['code'] == TemplateAnalyzer.CheckCode.SUCCESS:
            insert_points = check_info['data']['insert_points']
            total_points = sum(len(p_d) if isinstance(p_d, list) else 1 for p_d in insert_points.values())
            # print(f"模板校验成功，共有 {len(insert_points)} 个不同标签名，{total_points} 个标签实例")

            if show_detail:
                count = 1
                for p_n, p_d in insert_points.items():
                    if isinstance(p_d, list):
                        for instance in p_d:
                            # print(f'\t{count}、{p_n}：{instance["text"]}')
                            count += 1
                    else:
                        # print(f'\t{count}、{p_n}：{p_d["text"]}')
                        count += 1
        else:
            # print(f'模板校验失败\n\t错误代码：{check_info["code"]}\n\t错误信息：{check_info["msg"]}')
            pass
