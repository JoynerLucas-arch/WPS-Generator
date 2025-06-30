import time
from abc import ABCMeta, abstractmethod
from typing import Any, List
import os

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import CT_Tbl
from docx.table import Table, _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from helper import image_size
from helper.docx_helper import *
from helper.type_helper import *


class Label(metaclass=ABCMeta):
    """
    内容标签接口
    """

    @classmethod
    @abstractmethod
    def get_type(cls) -> str:
        """获取内容标签类型"""
        pass

    @classmethod
    @abstractmethod
    def has_content(cls) -> bool:
        """是否有插入内容（生成时外部输入的信息）"""
        pass

    @classmethod
    @abstractmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        """注册静态数据"""
        pass

    @classmethod
    @abstractmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        """在插入点插入数据"""
        pass

    @classmethod
    @abstractmethod
    def check_data_type(cls, data: Any) -> bool:
        """检查插入输入类型"""
        pass


class NoContentLabel(Label, metaclass=ABCMeta):
    @classmethod
    def has_content(cls) -> bool:
        return False

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        return True


class ContentLabel(Label, metaclass=ABCMeta):
    @classmethod
    def has_content(cls) -> bool:
        return True

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        pass


class LabelManager:
    __labels = []

    @classmethod
    def register(cls, label):
        if not issubclass(label, Label):
            raise TypeError('label must be sub class of type `Label`')
        cls.__labels.append(label)

    @classmethod
    def get_labels(cls) -> List[Label]:
        return cls.__labels.copy()

    @classmethod
    def print_registered_labels(cls):
        print(f'当前注册的内容标签类型有：{[l.get_type() for l in cls.__labels]}')
        print(f'无内容类型有：{[l.get_type() for l in cls.__labels if not l.has_content()]}')
        print(f'有内容类型有：{[l.get_type() for l in cls.__labels if l.has_content()]}')


class TextLabel(ContentLabel):
    """文本内容标签"""

    @classmethod
    def get_type(cls) -> str:
        return 'text'

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        # 处理跨run的情况
        if 'containing_runs' in point_data:
            containing_runs = point_data['containing_runs']

            # 在第一个run中放置完整的替换文本
            first_run = containing_runs[0]['run']
            first_run_start = containing_runs[0]['start']
            start_offset = point_data['start_pos'] - first_run_start

            # 获取标记之前的文本
            prefix = first_run.text[:start_offset]

            # 获取最后一个run中标记之后的文本
            last_run = containing_runs[-1]['run']
            last_run_start = containing_runs[-1]['start']
            end_offset = point_data['end_pos'] - last_run_start
            suffix = last_run.text[end_offset:]

            # 更新第一个run的文本
            first_run.text = prefix + data

            # 清空中间的run
            for run_info in containing_runs[1:-1]:
                run_info['run'].text = ''

            # 更新最后一个run的文本（如果它不是第一个run）
            if len(containing_runs) > 1:
                last_run.text = suffix
        else:
            # 向后兼容：如果没有containing_runs信息，使用原来的逻辑
            new_text = point_data['run'].text.replace(point_data['text'], data)
            point_data['run'].text = new_text

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        return isinstance(data, str)


LabelManager.register(TextLabel)


class DateLabel(NoContentLabel):
    @classmethod
    def get_type(cls) -> str:
        return 'date'

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        now_time = time.time()
        now_time_struct = time.localtime(now_time)
        date_s = time.strftime("%Y-%m-%d", now_time_struct)
        static_datas[cls.get_type()] = date_s

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        # 检查是否包含多个runs
        if 'containing_runs' in point_data:
            containing_runs = point_data['containing_runs']
            first_run = containing_runs[0]['run']
            first_run_start = containing_runs[0]['start']
            start_offset = point_data['start_pos'] - first_run_start
            prefix = first_run.text[:start_offset]

            last_run = containing_runs[-1]['run']
            last_run_start = containing_runs[-1]['start']
            end_offset = point_data['end_pos'] - last_run_start
            suffix = last_run.text[end_offset:]

            # 更新第一个run的文本
            first_run.text = prefix + static_datas[point_data['type']]

            # 清空中间的runs
            for run_info in containing_runs[1:-1]:
                run_info['run'].text = ''

            # 如果有多个runs，更新最后一个
            if len(containing_runs) > 1:
                last_run.text = suffix
        else:
            # 如果只有一个run，直接替换文本
            new_text = point_data['run'].text.replace(point_data['text'], static_datas[point_data['type']])
            point_data['run'].text = new_text

        # 如果是表格单元格中的标签，需要特殊处理段落格式
        if 'cell' in point_data:
            paragraph = point_data['paragraph']
            if hasattr(paragraph, 'paragraph_format'):
                # 保持段落的格式（如对齐方式等）
                paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment


LabelManager.register(DateLabel)


# In labels.py, update the TimeLabel class:

class TimeLabel(NoContentLabel):
    @classmethod
    def get_type(cls) -> str:
        return 'time'

    @classmethod
    def register_static_datas(cls, static_datas: dict) -> None:
        now_time = time.time()
        now_time_struct = time.localtime(now_time)
        time_s = time.strftime("%H:%M:%S", now_time_struct)
        static_datas[cls.get_type()] = time_s

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        # Check if containing_runs exists for handling multi-run cases
        if 'containing_runs' in point_data:
            containing_runs = point_data['containing_runs']
            first_run = containing_runs[0]['run']
            first_run_start = containing_runs[0]['start']
            start_offset = point_data['start_pos'] - first_run_start
            prefix = first_run.text[:start_offset]

            last_run = containing_runs[-1]['run']
            last_run_start = containing_runs[-1]['start']
            end_offset = point_data['end_pos'] - last_run_start
            suffix = last_run.text[end_offset:]

            first_run.text = prefix + static_datas[point_data['type']]

            for run_info in containing_runs[1:-1]:
                run_info['run'].text = ''

            if len(containing_runs) > 1:
                last_run.text = suffix
        else:
            new_text = point_data['run'].text.replace(point_data['text'], static_datas[point_data['type']])
            point_data['run'].text = new_text


LabelManager.register(TimeLabel)


class OrderedListLabel(ContentLabel):
    @classmethod
    def get_type(cls) -> str:
        return 'ordered-list'

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        paragraph = point_data['paragraph']
        for i, item in enumerate(data):
            p = paragraph.insert_paragraph_before(f'{i + 1}. {item}')
            copy_paragraph_style(p, paragraph)
        delete_paragraph(paragraph)

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data是list，且元素是str"""
        return isinstance(data, Iterable) and all([isinstance(d, str) for d in data])


LabelManager.register(OrderedListLabel)


class UnorderedListLabel(ContentLabel):
    _header_chars = {"circle0": "•", "square0": "▪", "disc0": "◦",
                     "circle1": "●", "square1": "■", "disc1": "○", "diamond1": "◆", "diamond1_e": "◇", }
    _default_header_char = "circle1"
    _default_header_gap = 1

    @classmethod
    def get_type(cls) -> str:
        return 'unordered-list'

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        paragraph = point_data['paragraph']
        for i, item in enumerate(data):
            p = paragraph.insert_paragraph_before(f'{cls._header_chars[cls._default_header_char]}{" " * cls._default_header_gap}{item}')
            copy_paragraph_style(p, paragraph)
        delete_paragraph(paragraph)

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data是list，且元素是str"""
        return isinstance(data, Iterable) and all([isinstance(d, str) for d in data])


LabelManager.register(UnorderedListLabel)


class ImageLabel(ContentLabel):
    @classmethod
    def get_type(cls) -> str:
        return 'image'

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        """支持多种表格形式中的图片插入，自适应大小居中显示，并且图片描述也居中"""
        try:
            pic_desc, pic_url = data
            document = point_data['document']
            paragraph = point_data['paragraph']
            
            # 获取当前段落所在的单元格和表格（如果在表格中）
            in_table = 'cell' in point_data
            cell = point_data.get('cell', None)
            
            # 增强图片路径处理
            import os
            if not os.path.exists(pic_url):
                # 尝试在模板文件所在目录查找
                template_dir = os.path.dirname(document.path) if hasattr(document, 'path') else ''
                possible_paths = [
                    pic_url,  # 原始路径
                    os.path.join(os.getcwd(), pic_url),  # 当前工作目录下
                    os.path.join(template_dir, pic_url),  # 模板文件所在目录
                    os.path.abspath(pic_url),  # 绝对路径
                    os.path.join(os.path.dirname(os.path.abspath(__file__)), pic_url),  # 脚本文件目录
                    os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', pic_url),  # 项目根目录
                ]
                
                # 检查所有可能的路径
                found = False
                for path in possible_paths:
                    if os.path.exists(path):
                        pic_url = path
                        found = True
                        break
                
                if not found:
                    print(f"错误：图片文件不存在 - {pic_url}")
                    # 如果图片不存在，保留标签
                    if 'run' in point_data and point_data['run'] is not None:
                        point_data['run'].text = point_data['text'] + "(找不到图片)"
                    return
            
            # 获取图片尺寸
            try:
                img_width, img_height = image_size.get(pic_url)
            except Exception as e:
                print(f"警告：获取图片尺寸时出错 - {str(e)}")
                # 使用合理的默认值
                img_width, img_height = 800, 600
            
            # 计算文档的基本尺寸
            d_section = document.sections[0]
            doc_width = d_section.page_width - d_section.left_margin - d_section.right_margin
            doc_height = d_section.page_height - d_section.top_margin - d_section.bottom_margin
            
            # 根据是否在表格中使用不同的尺寸计算策略
            if in_table:
                # 在表格中的图片尺寸计算
                try:
                    # 防止表格被图片撑大的关键措施：强制限制图片尺寸
                    
                    # 1. 获取单元格的实际宽度
                    from docx.shared import Inches, Cm
                    
                    # 尝试直接获取单元格宽度
                    try:
                        cell_width = cell.width
                    except:
                        # 如果无法直接获取，则使用估计值
                        cell_width = doc_width * 0.25  # 假设表格占据页面宽度的1/4
                    
                    # 2. 如果单元格宽度为None或不合理的值，使用安全默认值
                    if cell_width is None or cell_width <= 0:
                        cell_width = Cm(4)  # 使用4厘米作为默认宽度
                    
                    # 3. 设置图片最大宽度为单元格宽度的80%
                    max_width = cell_width * 0.8
                    
                    # 4. 设置图片最大高度
                    # 为描述文本预留空间
                    max_height = Cm(4) if pic_desc else Cm(5)  # 有描述则预留更多垂直空间
                    
                    # 5. 计算缩放比例
                    width_ratio = max_width / img_width
                    height_ratio = max_height / img_height
                    scale_ratio = min(width_ratio, height_ratio) * 0.95  # 额外的安全系数
                    
                    # 6. 计算最终图片尺寸
                    final_width = img_width * scale_ratio
                    
                    # 7. 尝试设置单元格属性以防止被图片撑大
                    try:
                        # 获取单元格的XML元素
                        tc = cell._tc
                        # 设置单元格不自动调整大小
                        from docx.oxml import parse_xml
                        tc_pr = tc.get_or_add_tcPr()
                        # 添加宽度固定属性
                        tc_pr.append(parse_xml(f'<w:tcW w:w="{int(cell_width)}" w:type="dxa"/>'))
                    except Exception as e:
                        print(f"警告：设置单元格固定宽度时出错 - {str(e)}")
                except Exception as e:
                    print(f"警告：计算表格中图片尺寸时出错 - {str(e)}")
                    # 使用安全的默认值
                    max_width = doc_width * 0.2
                    max_height = doc_height * 0.15
                    
                    width_ratio = max_width / img_width
                    height_ratio = max_height / img_height
                    scale_ratio = min(width_ratio, height_ratio)
                    final_width = img_width * scale_ratio
            else:
                # 不在表格中的图片处理（保持原有逻辑）
                max_width = doc_width * 0.8
                max_height = doc_height * 0.5
                
                width_ratio = max_width / img_width
                height_ratio = max_height / img_height
                scale_ratio = min(width_ratio, height_ratio)
                final_width = img_width * scale_ratio
            
            # 确定插入图片的段落
            if in_table:
                # 如果在表格中，我们替换标签所在段落
                ip = paragraph
                # 清空段落内容但保留格式
                for run in ip.runs:
                    run.text = ""
                ir = ip.add_run()
            else:
                # 否则，在当前段落前插入新段落
                ip = paragraph.insert_paragraph_before()
                ir = ip.add_run()
            
            # 插入图片
            try:
                picture = ir.add_picture(pic_url, width=final_width)
            except Exception as e:
                print(f"插入图片失败: {pic_url}, 错误: {str(e)}")
                # 在遇到错误时，保留原始标签
                if 'run' in point_data and point_data['run'] is not None:
                    point_data['run'].text = point_data['text'] + "(插入失败)"
                return
            
            # 设置段落居中对齐
            ip.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 处理描述文本
            if pic_desc is not None and pic_desc.strip():
                # 检查当前单元格中是否已经存在相同的图片描述
                desc_already_exists = False
                
                if in_table:
                    # 检查单元格中是否已经有相同的描述文本
                    for p in cell.paragraphs:
                        if p != ip and p.text.strip() == pic_desc.strip():
                            desc_already_exists = True
                            break
                
                # 如果描述文本已存在，则不添加
                if not desc_already_exists:
                    # 如果有描述文本，添加描述文本段落
                    if in_table:
                        # 如果在表格中，创建单元格中的新段落
                        desc_p = cell.add_paragraph()
                    else:
                        # 如果不在表格中，在原段落位置创建描述段落
                        desc_p = paragraph
                        # 清空原段落中所有run的文本
                        for run in desc_p.runs:
                            run.text = ""
                        
                    # 添加描述文本
                    desc_run = desc_p.add_run(pic_desc)
                    
                    # 如果有原始run，复制其样式
                    if len(paragraph.runs) > 0:
                        original_run = paragraph.runs[0]
                        copy_run_style(original_run, desc_run)
                        
                    # 设置描述段落居中对齐
                    desc_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                # 如果没有描述文本，并且不在表格中，则删除原段落
                if not in_table:
                    delete_paragraph(paragraph)
        except Exception as e:
            print(f"错误：插入图片时出现问题 - {str(e)}")
            import traceback
            traceback.print_exc()
            # 在遇到错误时，尝试保留原始标签
            if 'run' in point_data and point_data['run'] is not None:
                point_data['run'].text = point_data['text'] + "(处理出错)"

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data为tuple，第一个元素是描述字符串（可为None），第二个元素是图片url"""
        return isinstance(data, tuple) and len(data) == 2 and (
                    isinstance(data[0], str) or data[0] is None) and isinstance(data[1], str)

LabelManager.register(ImageLabel)


class LinkLabel(ContentLabel):
    @classmethod
    def get_type(cls) -> str:
        return 'link'

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        """将包含标签的 run 替换为 link"""
        link_n, link_url = data
        paragraph = point_data['paragraph']
        run_index = point_data['run_index']
        set_hyperlink(run_index, paragraph, link_n, link_url)

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """要求data是tuple，第一个元素是链接名称，第二个元素是链接url"""
        return isinstance(data, tuple) and len(data) == 2 and isinstance(data[0], str) and isinstance(data[1], str)


LabelManager.register(LinkLabel)


def set_cell_border(cell: _Cell, **kwargs):
    """
    设置单元格边框
    使用方法:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        start={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 检查是否存在tcBorders标签，如果没有则创建一个
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # 遍历所有可用边框位置
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # 检查边框标签是否存在，不存在则创建
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # 属性顺序很重要
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


class TableLabel(ContentLabel):
    @classmethod
    def get_type(cls) -> str:
        return 'table'

    @classmethod
    def insert_data_to_point(cls, point_data: dict, data: Any, static_datas: dict) -> None:
        """在内容标签的 paragraph 下插入表格，并删除内容标签的 paragraph"""
        document = point_data['document']
        paragraph = point_data['paragraph']
        document_body = document._body

        tbl = CT_Tbl.new_tbl(0, len(data[0]), document._block_width)
        paragraph._element.addnext(tbl)
        table = Table(tbl, document_body)

        for row in data:
            for i, cell in enumerate(table.add_row().cells):
                cell.text = row[i]
                # 设置单元格内容水平和垂直居中对齐（Word兼容）
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 设置表头、列头样式
        # for cell in table.row_cells(0):
        #     for r in cell.paragraphs[0].runs:
        #         r.bold = True
        # for cell in table.column_cells(0):
        #     for r in cell.paragraphs[0].runs:
        #         r.bold = False
                
        # 为所有单元格添加边框
        border_settings = {
            "top": {"sz": 6, "val": "single", "color": "#000000"},
            "bottom": {"sz": 6, "val": "single", "color": "#000000"},
            "start": {"sz": 6, "val": "single", "color": "#000000"},
            "end": {"sz": 6, "val": "single", "color": "#000000"}
        }
        
        # 为表格的每个单元格添加边框
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, **border_settings)

        delete_paragraph(paragraph)

    @classmethod
    def check_data_type(cls, data: Any) -> bool:
        """
        要求表格数据为二维矩阵 Iterable[Iterable[str]]，默认首行为表头，首列为列头，行/列数都不可为0
        """
        return isinstance(data, Iterable) and all(check_iterable_type(i, str) for i in data) and len(data) > 0 and len(data[0]) > 0


LabelManager.register(TableLabel)
